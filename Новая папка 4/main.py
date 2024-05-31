from docx import Document
from docx.shared import Pt

# Создание нового документа
doc = Document()
doc.add_heading('Приглашение на свадьбу', level=1)

# Добавление текста с различными стилями
p = doc.add_paragraph()
p.add_run('ФИО: ').bold = True
p.add_run('________________________\n')
p.add_run('Адрес: ').italic = True
p.add_run('________________________\n')
p.add_run('Время, дата: ').bold = True
p.add_run('________________________\n')

# Добавление изображения подписи
p = doc.add_paragraph()
p.add_run('Подпись: ').bold = True
run = p.add_run()
run.add_picture('pod.png', width=Pt(60))

# Добавление таблицы
table = doc.add_table(rows=2, cols=2)
table.cell(0, 0).text = 'Место:'
table.cell(0, 1).text = '_____________________________\n'
table.cell(1, 0).text = 'Печать:'
cell = table.cell(1, 1)
cell.text = ''
p = cell.paragraphs[0]
r = p.add_run()
r.add_picture('pech1.png', width=Pt(120))

# Добавление заголовка, параграфа и (не)нумерованного списка
doc.add_heading('Заголовок различного уровня', level=2)
doc.add_paragraph('Текст параграфа', style='Body Text')
doc.add_paragraph('Нумерованный список:', style='List Number')
doc.add_paragraph('Элемент списка 1', style='List Bullet')

# Сохранение документа
doc.save('test22.docx')
print("Создание документа завершено!")