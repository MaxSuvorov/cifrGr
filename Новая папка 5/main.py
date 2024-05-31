from docx import Document
from docx.shared import Pt

def generate_invitations(invitees_list, output_folder):
    for invitee in invitees_list:
        doc = Document()
        doc.add_heading('Приглашение на свадьбу', level=1)

        p = doc.add_paragraph()
        p.add_run('ФИО: ').bold = True
        p.add_run(invitee['name'] + '\n')
        p.add_run('Адрес: ').italic = True
        p.add_run(invitee['address'] + '\n')
        p.add_run('Время, дата: ').bold = True
        p.add_run(invitee['datetime'] + '\n')

        p = doc.add_paragraph()
        p.add_run('Подпись: ').bold = True
        run = p.add_run()
        run.add_picture('pod.png', width=Pt(60))

        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = 'Место:'
        table.cell(0, 1).text = '_____________________________\n'
        table.cell(1, 0).text = 'Печать:'
        cell = table.cell(1, 1)
        cell.text = ''
        p = cell.paragraphs[0]
        r = p.add_run()
        r.add_picture('pech1.png', width=Pt(120))

        doc.save(output_folder + '/' + invitee['name'] + '.docx')

# Пример списка приглашенных
invitees = [
    {'name': 'Иванов Иван Иванович', 'address': 'ул. Пушкина д.10', 'datetime': '12 апреля 2023, 17:00'},
    {'name': 'Пупкин Пуп Пупович', 'address': 'ул. Пушкина д.Калатушкина', 'datetime': '2 июня 2024, 9:00'},
    {'name': 'Сигизмундов Антон Никитич', 'address': 'ул. Кефира д.12', 'datetime': '3 февраля 2021, 10:00'},
    {'name': 'Пугачева Бритни Меладзовна', 'address': 'ул. Мира д.10', 'datetime': '8 мая 2020, 15:00'},
    {'name': 'Петрова Анна Сергеевна', 'address': 'пр. Мира д.5', 'datetime': '15 апреля 2023, 12:00'}
]

# Вызов функции для генерации документов с приглашениями
generate_invitations(invitees, 'output_folder')

print("Создание документов в папке output_folder завершено")